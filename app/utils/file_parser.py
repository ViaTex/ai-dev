"""
Utility functions and router for parsing resume files (PDF and DOCX).

Supports:
- Native PDF text extraction
- Scanned PDF OCR
- DOCX text extraction
"""

# -----------------------------
# Imports
# -----------------------------
import io
import os
from typing import Optional

from fastapi import UploadFile, HTTPException

import pdfplumber
import fitz  # PyMuPDF
from pdf2image import convert_from_path
import pytesseract

import docx
from docx import Document


# ============================================================
# SECTION 1: ASYNC FILE PARSERS (FastAPI UploadFile)
# ============================================================

async def extract_text_from_pdf(file: UploadFile) -> str:
    """
    Extract text from an uploaded PDF using pdfplumber.
    """
    try:
        content = await file.read()
        text = ""

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="No text could be extracted from the PDF"
            )

        return text.strip()

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error parsing PDF: {str(e)}"
        )


async def extract_text_from_docx(file: UploadFile) -> str:
    """
    Extract text from an uploaded DOCX file.
    """
    try:
        content = await file.read()
        doc = Document(io.BytesIO(content))

        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text

        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="No text could be extracted from the DOCX"
            )

        return text.strip()

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error parsing DOCX: {str(e)}"
        )


async def parse_resume_file(file: UploadFile) -> str:
    """
    Entry point for FastAPI resume parsing.
    Routes file to the correct extractor.
    """
    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        return await extract_text_from_pdf(file)

    elif filename.endswith(".docx"):
        return await extract_text_from_docx(file)

    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Only PDF and DOCX are supported."
        )


# ============================================================
# SECTION 2: FILE PATH BASED ROUTER (OCR + Native PDF)
# ============================================================

class ResumeRouter:
    """
    Routes local resume files to the correct parsing pipeline.
    """

    def __init__(self, tesseract_cmd: Optional[str] = None):
        # Optional manual Tesseract path configuration
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    # -----------------------------
    # Main Router
    # -----------------------------
    def route_and_parse(self, file_path: str) -> str:
        """
        Detect file type and extract text accordingly.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file {file_path} does not exist.")

        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext in [".docx", ".doc"]:
            print("[LOG] DOCX detected → DOCX parser")
            return self._parse_docx(file_path)

        elif ext == ".pdf":
            if self._is_scanned_pdf(file_path):
                print("[LOG] Scanned PDF detected → OCR pipeline")
                return self._pipeline_ocr(file_path)
            else:
                print("[LOG] Native PDF detected → Text extraction pipeline")
                return self._pipeline_native_pdf(file_path)

        else:
            raise ValueError(f"Unsupported file format: {ext}")

    # -----------------------------
    # PDF Type Detection
    # -----------------------------
    def _is_scanned_pdf(self, pdf_path: str) -> bool:
        """
        Detect whether a PDF is scanned (image-based).
        """
        try:
            with fitz.open(pdf_path) as doc:
                if len(doc) == 0:
                    return True

                first_page_text = doc[0].get_text().strip()
                return len(first_page_text) < 5

        except Exception as e:
            print(f"[ERROR] PDF scan check failed: {e}")
            return True

    # -----------------------------
    # Native PDF Pipeline
    # -----------------------------
    def _pipeline_native_pdf(self, pdf_path: str) -> str:
        """
        Extract text from a text-based PDF.
        """
        text_content = []

        with fitz.open(pdf_path) as doc:
            for page in doc:
                text_content.append(page.get_text())

        return "\n".join(text_content)

    # -----------------------------
    # OCR Pipeline
    # -----------------------------
    def _pipeline_ocr(self, pdf_path: str) -> str:
        """
        Extract text from scanned PDFs using OCR.
        """
        text_content = []
        images = convert_from_path(pdf_path)

        for idx, image in enumerate(images, start=1):
            print(f"[LOG] OCR processing page {idx}")
            text = pytesseract.image_to_string(image)
            text_content.append(text)

        return "\n".join(text_content)

    # -----------------------------
    # DOCX Parser
    # -----------------------------
    def _parse_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file including tables.
        """
        doc = docx.Document(file_path)
        content = []

        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if row_text:
                    content.append(" | ".join(row_text))

        return "\n".join(content)
