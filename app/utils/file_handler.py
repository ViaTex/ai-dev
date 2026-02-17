from __future__ import annotations

import io
import re
from typing import Dict, List, Tuple

import anyio
import pdfplumber
from docx import Document
from docx.oxml import CT_Hyperlink
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from fastapi import UploadFile

from app.utils.validators import FileTooLargeError, ParsingError


async def read_upload_file(upload_file: UploadFile, max_size_mb: int) -> bytes:
    max_bytes = max_size_mb * 1024 * 1024
    buffer = bytearray()
    while True:
        chunk = await upload_file.read(1024 * 1024)
        if not chunk:
            break
        buffer.extend(chunk)
        if len(buffer) > max_bytes:
            raise FileTooLargeError("File exceeds maximum size")
    await upload_file.close()
    return bytes(buffer)


def _extract_pdf_hyperlinks(file_bytes: bytes) -> List[Dict[str, str]]:
    """Extract hyperlinks from PDF file."""
    hyperlinks = []
    
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                # Extract annotations (clickable links)
                if hasattr(page, 'annots') and page.annots:
                    for annot in page.annots:
                        if annot and isinstance(annot, dict):
                            # Get the URI from the annotation
                            uri = annot.get('uri') or annot.get('A', {}).get('URI')
                            # Get the display text from the rectangle area
                            text = annot.get('contents') or ""
                            
                            if uri:
                                hyperlinks.append({
                                    'text': text.strip(),
                                    'url': uri.strip()
                                })
                
                # Also check for hyperlinks in page objects
                if hasattr(page, 'hyperlinks'):
                    for link in page.hyperlinks:
                        if isinstance(link, dict):
                            url = link.get('uri') or link.get('url')
                            text = link.get('text', '')
                            if url:
                                hyperlinks.append({
                                    'text': text.strip(),
                                    'url': url.strip()
                                })
    except Exception:
        # If hyperlink extraction fails, return empty list
        pass
    
    return hyperlinks


def _extract_docx_hyperlinks(file_bytes: bytes) -> List[Dict[str, str]]:
    """Extract hyperlinks from DOCX file."""
    hyperlinks = []
    
    try:
        document = Document(io.BytesIO(file_bytes))
        
        # Iterate through all paragraphs
        for paragraph in document.paragraphs:
            # Get hyperlinks from paragraph XML
            for child in paragraph._element.iterchildren():
                if isinstance(child, CT_Hyperlink):
                    # Get hyperlink text
                    text = ''.join(node.text for node in child.iter() if hasattr(node, 'text') and node.text)
                    
                    # Get hyperlink URL
                    # Hyperlinks can have rId attribute that references a relationship
                    r_id = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if r_id:
                        try:
                            url = document.part.rels[r_id].target_ref
                            hyperlinks.append({
                                'text': text.strip(),
                                'url': url.strip()
                            })
                        except (KeyError, AttributeError):
                            pass
        
        # Also check hyperlinks in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for child in paragraph._element.iterchildren():
                            if isinstance(child, CT_Hyperlink):
                                text = ''.join(node.text for node in child.iter() if hasattr(node, 'text') and node.text)
                                r_id = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                if r_id:
                                    try:
                                        url = document.part.rels[r_id].target_ref
                                        hyperlinks.append({
                                            'text': text.strip(),
                                            'url': url.strip()
                                        })
                                    except (KeyError, AttributeError):
                                        pass
    except Exception:
        # If hyperlink extraction fails, return empty list
        pass
    
    return hyperlinks


def _extract_pdf_text(file_bytes: bytes) -> str:
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages_text = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages_text).strip()


def _extract_docx_text(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return "\n".join(paragraphs).strip()


def _extract_pdf_with_links(file_bytes: bytes) -> Tuple[str, List[Dict[str, str]]]:
    """Extract both text and hyperlinks from PDF."""
    text = _extract_pdf_text(file_bytes)
    hyperlinks = _extract_pdf_hyperlinks(file_bytes)
    return text, hyperlinks


def _extract_docx_with_links(file_bytes: bytes) -> Tuple[str, List[Dict[str, str]]]:
    """Extract both text and hyperlinks from DOCX."""
    text = _extract_docx_text(file_bytes)
    hyperlinks = _extract_docx_hyperlinks(file_bytes)
    return text, hyperlinks


async def extract_text(file_bytes: bytes, file_extension: str) -> str:
    if file_extension == "pdf":
        return await anyio.to_thread.run_sync(_extract_pdf_text, file_bytes)
    if file_extension == "docx":
        return await anyio.to_thread.run_sync(_extract_docx_text, file_bytes)
    raise ParsingError("Unsupported file type for extraction")


async def extract_text_with_links(file_bytes: bytes, file_extension: str) -> Tuple[str, List[Dict[str, str]]]:
    """Extract both text and hyperlinks from the file."""
    if file_extension == "pdf":
        return await anyio.to_thread.run_sync(_extract_pdf_with_links, file_bytes)
    if file_extension == "docx":
        return await anyio.to_thread.run_sync(_extract_docx_with_links, file_bytes)
    raise ParsingError("Unsupported file type for extraction")
